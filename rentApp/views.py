from django.shortcuts import render
from rest_framework import viewsets, permissions, status
from rest_framework.decorators import api_view, permission_classes, action
from rest_framework.response import Response
from rest_framework.views import APIView
from rest_framework.authtoken.models import Token
from django.contrib.auth import authenticate
from django.db.models import Sum, Q
from django.utils import timezone
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
import io
from rest_framework_simplejwt.tokens import RefreshToken
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, RGBColor
from urllib.parse import quote
import os
from decimal import Decimal
from rest_framework.permissions import IsAuthenticated, AllowAny

from .models import Role, User, Car, Rental, Maintenance, Penalty, Discount
from .serializers import (
    RoleSerializer, UserSerializer, CarSerializer, RentalSerializer,
    MaintenanceSerializer, PenaltySerializer, DiscountSerializer,
    RentalCreateSerializer, RentalOperatorSerializer, UserRegistrationSerializer
)
from .permissions import IsOperator

# Create your views here.

class RoleViewSet(viewsets.ModelViewSet):
    queryset = Role.objects.all()
    serializer_class = RoleSerializer
    permission_classes = [permissions.IsAdminUser]

class UserViewSet(viewsets.ModelViewSet):
    queryset = User.objects.all()
    serializer_class = UserSerializer
    
    def get_permissions(self):
        if self.action == 'create':
            return [permissions.AllowAny()]
        return [permissions.IsAdminUser()]

class CarViewSet(viewsets.ModelViewSet):
    queryset = Car.objects.all()
    serializer_class = CarSerializer
    
    @action(detail=False, methods=['get'])
    def available(self, request):
        """Получить список доступных автомобилей"""
        cars = Car.objects.filter(status='available')
        serializer = self.get_serializer(cars, many=True)
        return Response(serializer.data)
    
    @action(detail=True, methods=['post'])
    def maintenance(self, request, pk=None):
        """Отправить автомобиль на техническое обслуживание"""
        car = self.get_object()
        
        # Проверяем, что автомобиль доступен
        if car.status != 'available':
            return Response(
                {'error': 'Автомобиль не доступен для обслуживания'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Создаем запись о техническом обслуживании
        maintenance = Maintenance.objects.create(
            car=car,
            description=request.data.get('description', 'Плановое обслуживание'),
            status='pending'
        )
        
        # Обновляем статус автомобиля
        car.status = 'maintenance'
        car.save()
        
        serializer = MaintenanceSerializer(maintenance)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def financial_history(self, request):
        """Получить историю доходов и расходов по каждой машине"""
        try:
            # Получаем все машины
            cars = Car.objects.all()
            
            # Создаем словарь для хранения финансовой информации по каждой машине
            car_finances = {}
            
            # Для каждой машины получаем историю аренд и обслуживаний
            for car in cars:
                car_id = car.id
                
                # Получаем завершенные аренды для этой машины
                rentals = Rental.objects.filter(
                    car=car,
                    status='completed'
                )
                
                # Рассчитываем общий доход от аренд
                total_income = rentals.aggregate(Sum('total_price'))['total_price__sum'] or 0
                
                # Получаем завершенные обслуживания для этой машины
                maintenances = Maintenance.objects.filter(
                    car=car,
                    status='completed'
                )
                
                # Рассчитываем общие расходы на обслуживание
                total_expenses = maintenances.aggregate(Sum('cost'))['cost__sum'] or 0
                
                # Рассчитываем эффективность (прибыльность)
                efficiency = 0
                if total_income > 0:
                    efficiency = round((total_income - total_expenses) / total_income * 100)
                
                # Сохраняем финансовую информацию для этой машины
                car_finances[car_id] = {
                    'id': car_id,
                    'brand': car.brand,
                    'model': car.model,
                    'total_income': float(total_income),
                    'total_expenses': float(total_expenses),
                    'efficiency': efficiency
                }
            
            # Преобразуем словарь в список для ответа API
            result = list(car_finances.values())
            
            return Response(result)
        except Exception as e:
            print(f"Ошибка при получении финансовой истории: {str(e)}")
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class RentalViewSet(viewsets.ModelViewSet):
    queryset = Rental.objects.all()
    permission_classes = [permissions.IsAuthenticated]
    
    def get_serializer_class(self):
        if self.action == 'create':
            return RentalCreateSerializer
        return RentalSerializer

    def get_queryset(self):
        return Rental.objects.filter(user=self.request.user)

    def create(self, request, *args, **kwargs):
        
        try:
            car_id = request.data.get('car_id')
            car = Car.objects.get(id=car_id)
            
            if car.status == 'in_rent':
                return Response(
                    {'error': 'Автомобиль уже в аренде'},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Логируем информацию о скидке
            applied_discount = request.data.get('applied_discount', 0)
            
            serializer = self.get_serializer(data=request.data)
            if serializer.is_valid():
                rental = serializer.save()
                # Устанавливаем статус автомобиля "ожидает подтверждения"
                car.status = 'pending'
                car.save()
                
                
                return Response(
                    RentalSerializer(rental).data,
                    status=status.HTTP_201_CREATED
                )
            return Response(
                serializer.errors,
                status=status.HTTP_400_BAD_REQUEST
            )
            
        except Car.DoesNotExist:
            return Response(
                {'error': 'Автомобиль не найден'},
                status=status.HTTP_404_NOT_FOUND
            )

    @action(detail=True, methods=['post'])
    def return_car(self, request, pk=None):
        rental = self.get_object()
        if rental.status != 'active':
            return Response(
                {'error': 'Можно завершать только активные аренды'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Обновляем статус аренды на "завершена"
        rental.status = 'completed'
        rental.return_date = timezone.now()  # Используем timezone.now() для правильного формата
        rental.return_condition = request.data.get('return_condition', '')
        rental.save()
        
        # Логируем для отладки
        
        # Обновляем статус автомобиля на "доступен"
        car = rental.car
        car.status = 'available'
        
        # Обновляем состояние автомобиля в зависимости от повреждений
        damage_level = request.data.get('damage_level', None)
        
        # Определяем новое состояние на основе повреждений
        new_condition = None
        if damage_level == 'minor':
            new_condition = 'good'
        elif damage_level == 'medium':
            new_condition = 'satisfactory'
        elif damage_level == 'severe':
            new_condition = 'needs_repair'
        
        # Обновляем состояние только если новое повреждение серьезнее предыдущего
        if new_condition:
            # Определяем приоритет состояний (чем выше число, тем хуже состояние)
            condition_priority = {
                'excellent': 1,
                'good': 2,
                'satisfactory': 3,
                'needs_repair': 4
            }
            
            current_priority = condition_priority.get(car.condition, 0)
            new_priority = condition_priority.get(new_condition, 0)
            
            # Обновляем состояние только если новое состояние хуже текущего
            if new_priority > current_priority:
                car.condition = new_condition
        
        car.save()
        
        # Рассчитываем штраф на основе уровня топлива и повреждений
        try:
            fuel_level = int(request.data.get('fuel_level', 100))
            # Ограничиваем значение от 0 до 100
            fuel_level = max(0, min(100, fuel_level))
        except (ValueError, TypeError):
            # Если fuel_level не может быть преобразован в число, используем значение по умолчанию
            fuel_level = 100
        
        damage_level_russian = request.data.get('damage_level_russian', 'Нет')
        
        penalty_amount = 0
        penalty_description = []
        
        # Штраф за низкий уровень топлива
        if fuel_level < 50:
            fuel_penalty = Decimal('5000')  # Фиксированный штраф за низкий уровень топлива
            penalty_amount += fuel_penalty
            penalty_description.append(f"Низкий уровень топлива ({fuel_level}%): {fuel_penalty} руб.")
        
        # Штраф за повреждения
        if damage_level:
            damage_penalty = 0
            if damage_level == 'minor':
                damage_penalty = rental.total_price * Decimal('0.5')
            elif damage_level == 'medium':
                damage_penalty = rental.total_price * Decimal('1.0')
            elif damage_level == 'severe':
                damage_penalty = rental.total_price * Decimal('1.5')
            
            if damage_penalty > 0:
                penalty_amount += damage_penalty
                penalty_description.append(f"Повреждения ({damage_level_russian}): {damage_penalty} руб.")
        
        # Создаем штраф, если есть
        if penalty_amount > 0:
            Penalty.objects.create(
                rental=rental,
                amount=penalty_amount,
                description="; ".join(penalty_description)
            )
        
        return Response({'status': 'success'})

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_agreement(request):
    try:
        car_id = request.data.get('car_id')
        start_date = request.data.get('start_date')
        end_date = request.data.get('end_date')
        personal_info = request.data.get('personal_info')
        total_price = request.data.get('total_price')
        
        car = Car.objects.get(id=car_id)
        current_date = datetime.now().strftime('%d.%m.%Y')
        
        
        doc = Document()
        
        # Установка стиля для всего документа
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)  # Основной текст 14pt
        
        # Заголовок
        heading = doc.add_paragraph('ДОГОВОР АРЕНДЫ АВТОМОБИЛЯ')
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.runs[0].font.size = Pt(14)
        heading.paragraph_format.space_after = Pt(12)
        
        # Место и дата
        date = doc.add_paragraph(f'г. Санкт-Петербург\t\t\t\t\t{current_date}')
        date.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Выравнивание по ширине
        date.paragraph_format.space_after = Pt(12)
        
        # Преамбула
        preamble = doc.add_paragraph()
        preamble.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Выравнивание по ширине
        preamble.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        preamble.add_run('ООО "Sewxrr RentCar" в лице генерального директора sewxrr, действующего на основании Устава, именуемый в дальнейшем «Арендодатель», с одной стороны, и гр. ')
        preamble.add_run(f'{personal_info["fullName"]}').bold = True
        preamble.add_run(', паспорт: серия ')
        preamble.add_run(f'{personal_info["passportNumber"]}').bold = True
        preamble.add_run(f', проживающий по адресу: {personal_info["address"]}, именуемый в дальнейшем «Арендатор», с другой стороны, именуемые в дальнейшем «Стороны», заключили настоящий договор, в дальнейшем «Договор», о нижеследующем:')
        preamble.paragraph_format.space_after = Pt(30)
        
        # Функция для добавления заголовков разделов
        def add_section_heading(text):
            heading = doc.add_paragraph(text)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading.runs[0].font.size = Pt(16)  # Заголовки разделов 16pt
            heading.runs[0].font.bold = False   # Убираем жирный шрифт
            heading.paragraph_format.space_before = Pt(30)
            heading.paragraph_format.space_after = Pt(12)
            return heading
        
        # Функция для добавления параграфа с нужным форматированием
        def add_formatted_paragraph(text, space_after=12):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.add_run(text)
            p.paragraph_format.space_after = Pt(space_after)
            return p
        
        # 1. Предмет договора
        add_section_heading('1. ПРЕДМЕТ ДОГОВОРА')
        add_formatted_paragraph('1.1. Арендодатель предоставляет Арендатору следующее транспортное средство:\n\n' +
                              f'легковой автомобиль марка {car.brand} {car.model}, год выпуска {car.year} ' +
                              '(далее - Автомобиль), во временное владение и пользование за плату, а также оказывает ' +
                              'Арендатору своими силами услуги по управлению автомобилем и его технической эксплуатации.',
                              space_after=30)
        
        # 2. Условия договора
        add_section_heading('2. УСЛОВИЯ ДОГОВОРА')
        
        conditions = [
            'Арендодатель предоставляет автомобиль в исправном состоянии по Акту приема-передачи, являющемся неотъемлемой частью настоящего договора.',
            'Арендатор обязуется по истечение срока действия договора вернуть автомобиль в состоянии соответствующем отраженному в Акте приема-передачи, с учетом нормального износа.',
            'Арендатор производит текущий ремонт автомобиля за свой счет.',
            'Арендодателю предоставляется право использовать в нерабочее время сданный в аренду автомобиль в личных целях, с употреблением собственных горюче-смазочных материалов (бензин и т.п.).',
            'При использовании автомобиля в соответствии с п.2.4 стороны обязаны передавать автомобиль друг другу в исправном состоянии.'
        ]
        
        for i, text in enumerate(conditions, 1):
            add_formatted_paragraph(f'2.{i}. {text}')
        
        # 3. Порядок расчетов
        add_section_heading('3. ПОРЯДОК РАСЧЕТОВ')
        
        add_formatted_paragraph(f'3.1. Арендатор обязуется заплатить за аренду автомобиля {total_price} рублей.',
                              space_after=30)
        
        # 4. Срок действия договора
        add_section_heading('4. СРОК ДЕЙСТВИЯ ДОГОВОРА')
        
        add_formatted_paragraph(f'4.1. Договор заключен на срок с {start_date} по {end_date} и может быть продлен ' +
                              'сторонами по взаимному соглашению.',
                              space_after=30)
        
        # 5. Ответственность сторон
        add_section_heading('5. ОТВЕТСТВЕННОСТЬ СТОРОН')
        
        add_formatted_paragraph('5.1. Арендатор несет ответственность за сохранность арендуемого автомобиля в ' +
                              'рабочее время и в случае утраты или повреждения автомобиля в это время обязан ' +
                              'возместить Арендодателю причиненный ущерб, либо предоставить равноценный автомобиль ' +
                              'в течение 5 дней после его утраты или повреждения. В случае задержки возмещения ' +
                              'ущерба либо предоставления равноценного автомобиля в указанный срок, Арендатор ' +
                              'уплачивает пеню в размере 0.1% от стоимости ущерба либо оценочной стоимости автомобиля.')
        
        add_formatted_paragraph('5.2. Ответственность за сохранность автомобиля в нерабочее время несет ' +
                              'Арендодатель. При повреждении или утрате сданного в аренду автомобиля при ' +
                              'использовании в соответствии с п.2.3 настоящего договора Арендодатель обязан ' +
                              'устранить повреждения за свой счет или возместить Арендатору причиненный убыток. ' +
                              'Размер возмещения определяется соглашением сторон.',
                              space_after=30)
        
        # 6. Другие условия
        add_section_heading('6. ДРУГИЕ УСЛОВИЯ')
        for i, text in enumerate(conditions, 1):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.add_run(f'6.{i}. {text}')
            if i == len(conditions):  # Если это последний подпункт
                p.paragraph_format.space_after = Pt(24)  # Большой отступ после последнего подпункта
            else:
                p.paragraph_format.space_after = Pt(12)  # Убираем отступы между подпунктами
        
        # 7. Юридические адреса и реквизиты сторон
        add_section_heading('7. ЮРИДИЧЕСКИЕ АДРЕСА И РЕКВИЗИТЫ СТОРОН')
        add_formatted_paragraph('Арендодатель:')
        add_formatted_paragraph('ООО "Sewxrr RentCar"')
        add_formatted_paragraph('Адрес: г. Санкт-Петербург, ул. Примерная, д. 1')
        add_formatted_paragraph('ИНН/КПП: 1234567890/123456789')
        add_formatted_paragraph('р/с: 40702810123450123456', space_after=12)
        
        # Арендатор
        add_formatted_paragraph('Арендатор:')
        add_formatted_paragraph(f'ФИО: {personal_info["fullName"]}')
        add_formatted_paragraph(f'Паспорт: {personal_info["passportNumber"]}')
        add_formatted_paragraph(f'Адрес: {personal_info["address"]}')
        add_formatted_paragraph(f'Телефон: {personal_info["phone"]}')
        add_formatted_paragraph(f'Email: {personal_info["email"]}', space_after=30)
        
        # 8. Подписи сторон
        add_section_heading('8. ПОДПИСИ СТОРОН')
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        cell1 = table.cell(0, 0)
        cell1.text = 'Арендодатель:\n\nООО "Sewxrr RentCar"\n\n_____________ /___________/'
        
        cell2 = table.cell(0, 1)
        cell2.text = f'Арендатор:\n\n{personal_info["fullName"]}\n\n_____________ /___________/'
        
        # Сохранение документа
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Создание response с правильным именем файла
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Добавляем заголовки CORS
        response['Access-Control-Allow-Origin'] = '*'
        response['Access-Control-Allow-Methods'] = 'GET, OPTIONS'
        response['Access-Control-Allow-Headers'] = 'Content-Type, Authorization'
        
        # Кодируем имя файла для корректного отображения кириллицы
        filename = 'Договор аренды автомобиля.docx'
        response['Content-Disposition'] = f'attachment; filename="{quote(filename)}"'
        return response
        
    except Exception as e:
        print(f"Error generating agreement: {str(e)}")
        return Response({'error': str(e)}, status=400)

class MaintenanceViewSet(viewsets.ModelViewSet):
    queryset = Maintenance.objects.all()
    serializer_class = MaintenanceSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    @action(detail=False, methods=['get'])
    def cars(self, request):
        """Получить список автомобилей на обслуживании"""
        # Получаем все активные записи о техническом обслуживании
        maintenances = Maintenance.objects.filter(
            car__status='maintenance',
            status__in=['pending', 'in_progress']
        ).select_related('car')
        
        # Сериализуем данные о техническом обслуживании
        serializer = self.get_serializer(maintenances, many=True)
        return Response(serializer.data)
    
    @action(detail=True, methods=['post'])
    def accept(self, request, pk=None):
        """Принять заявку на обслуживание в работу"""
        maintenance = self.get_object()
        
        # Проверяем, что заявка в статусе ожидания
        if maintenance.status != 'pending':
            return Response(
                {'error': 'Можно принять в работу только заявки в статусе "В ожидании"'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Обновляем статус заявки
        maintenance.status = 'in_progress'
        maintenance.save()
        
        # Возвращаем обновленные данные о техническом обслуживании
        serializer = self.get_serializer(maintenance)
        return Response(serializer.data)
    
    @action(detail=True, methods=['patch'])
    def complete(self, request, pk=None):
        """Завершить техническое обслуживание"""
        maintenance = self.get_object()
        car = maintenance.car
        
        # Проверяем, что автомобиль на обслуживании
        if car.status != 'maintenance':
            return Response(
                {'error': 'Автомобиль не находится на обслуживании'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Проверяем, что заявка в работе
        if maintenance.status != 'in_progress':
            return Response(
                {'error': 'Можно завершить только заявки в статусе "В работе"'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Обновляем данные о техническом обслуживании
        description = request.data.get('description', '')
        cost = request.data.get('cost', 0)
        
        # Проверяем, что стоимость не отрицательная
        try:
            cost_value = float(cost)
            if cost_value < 0:
                return Response(
                    {'error': 'Стоимость ремонта не может быть отрицательной'},
                    status=status.HTTP_400_BAD_REQUEST
                )
        except (ValueError, TypeError):
            return Response(
                {'error': 'Некорректное значение стоимости ремонта'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        maintenance.description = description
        maintenance.cost = cost
        maintenance.status = 'completed'
        maintenance.completed_date = timezone.now().date()
        maintenance.save()
        
        # Обновляем статус и состояние автомобиля
        car.status = 'available'
        car.condition = 'excellent'  # После обслуживания состояние становится отличным
        car.save()
        
        # Возвращаем обновленные данные о техническом обслуживании
        serializer = self.get_serializer(maintenance)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def completed(self, request):
        """Получить список автомобилей с датой последнего обслуживания"""
        # Получаем список уникальных автомобилей, которые проходили обслуживание
        cars_with_maintenance = Car.objects.filter(
            maintenance__status='completed'
        ).distinct()
        
        # Для каждого автомобиля находим последнее обслуживание
        result = []
        for car in cars_with_maintenance:
            last_maintenance = Maintenance.objects.filter(
                car=car,
                status='completed'
            ).order_by('-completed_date').first()
            
            if last_maintenance:
                car_data = CarSerializer(car).data
                car_data['last_maintenance_date'] = last_maintenance.completed_date
                car_data['last_maintenance_id'] = last_maintenance.id
                result.append(car_data)
        
        return Response(result)
    
    @action(detail=True, methods=['get'])
    def history(self, request, pk=None):
        """Получить историю обслуживания для конкретного автомобиля"""
        try:
            car = Car.objects.get(pk=pk)
            maintenances = Maintenance.objects.filter(car=car, status='completed')
            serializer = self.get_serializer(maintenances, many=True)
            return Response(serializer.data)
        except Car.DoesNotExist:
            return Response(
                {'error': 'Автомобиль не найден'},
                status=status.HTTP_404_NOT_FOUND
            )

class PenaltyViewSet(viewsets.ModelViewSet):
    queryset = Penalty.objects.all()
    serializer_class = PenaltySerializer
    permission_classes = [permissions.IsAdminUser]

class DiscountViewSet(viewsets.ModelViewSet):
    queryset = Discount.objects.all()
    serializer_class = DiscountSerializer
    permission_classes = [permissions.IsAdminUser]

class UserPenaltyViewSet(viewsets.ReadOnlyModelViewSet):
    serializer_class = PenaltySerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return Penalty.objects.filter(rental__user=self.request.user)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def user_penalties(request):
    try:
        # Получаем штрафы через связь с арендой
        penalties = Penalty.objects.filter(rental__user=request.user)
        serializer = PenaltySerializer(penalties, many=True)
        return Response(serializer.data)
    except Exception as e:
        # Добавим логирование для диагностики ошибки
        print(f"Ошибка при получении штрафов: {str(e)}")
        return Response(
            {'error': f'Внутренняя ошибка сервера: {str(e)}'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )

@api_view(['POST'])
@permission_classes([AllowAny])
def register(request):
    print("Received registration data:", request.data)
    serializer = UserRegistrationSerializer(data=request.data)
    if serializer.is_valid():
        try:
            user = serializer.save()
            refresh = RefreshToken.for_user(user)
            return Response({
                'access': str(refresh.access_token),
                'refresh': str(refresh),
                'role': 'client'
            }, status=status.HTTP_201_CREATED)
        except Exception as e:
            print("Error during registration:", str(e))
            return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)
    print("Serializer errors:", serializer.errors)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

@api_view(['GET', 'PUT'])
@permission_classes([IsAuthenticated])
def user_profile(request):
    user = request.user  # Получаем текущего пользователя

    if request.method == 'GET':
        serializer = UserSerializer(user)  # Используем сериализатор для User
        return Response(serializer.data)

    elif request.method == 'PUT':
        serializer = UserSerializer(user, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def user_rentals(request):
    rentals = Rental.objects.filter(user=request.user)
    serializer = RentalSerializer(rentals, many=True)
    return Response(serializer.data)

class LoginView(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        username = request.data.get('username')
        password = request.data.get('password')
        user = authenticate(username=username, password=password)
        
        if user:
            token, _ = Token.objects.get_or_create(user=user)
            role = user.role.name if hasattr(user, 'role') else None
            
            return Response({
                'token': str(token.key),
                'user': {
                    'id': user.id,
                    'username': user.username,
                    'role': role,
                    'profile': {
                        'full_name': user.get_full_name(),
                        'email': user.email,
                        'phone': user.phone,
                        'address': user.address,
                        'passport_number': user.passport_number,
                        'driver_license': user.driver_license
                    }
                }
            })
        return Response(
            {'error': 'Неверные учетные данные'}, 
            status=status.HTTP_400_BAD_REQUEST
        )

class RegisterView(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        serializer = UserRegistrationSerializer(data=request.data)
        if serializer.is_valid():
            user = serializer.save()
            token, _ = Token.objects.get_or_create(user=user)
            return Response({
                'token': token.key,
                'user_id': user.id,
                'username': user.username
            }, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class ProfileView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        # Используем напрямую объект User вместо Profile
        serializer = UserSerializer(request.user)
        return Response(serializer.data)

    def put(self, request):
        # Обновляем данные непосредственно в модели User
        serializer = UserSerializer(request.user, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class OperatorRentalViewSet(viewsets.ModelViewSet):
    serializer_class = RentalOperatorSerializer
    permission_classes = [IsAuthenticated, IsOperator]
    
    def get_queryset(self):
        status_filter = self.request.query_params.get('status', None)
        queryset = Rental.objects.all().order_by('-created_at')
        if status_filter:
            queryset = queryset.filter(status=status_filter)
        return queryset

    @action(detail=True, methods=['post'])
    def approve(self, request, pk=None):
        rental = self.get_object()
        if rental.status != 'pending':
            return Response(
                {'error': 'Можно подтверждать только заявки в статусе "Ожидает подтверждения"'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        rental.status = 'active'
        rental.approved_by = request.user
        rental.approved_at = timezone.now()
        rental.save()
        
        # Обновляем статус автомобиля на "в аренде"
        car = rental.car
        car.status = 'in_rent'
        car.save()
        
        return Response(RentalOperatorSerializer(rental).data)

    @action(detail=True, methods=['post'])
    def complete_return(self, request, pk=None):
        rental = self.get_object()
        if rental.status != 'active':
            return Response(
                {'error': 'Можно завершать только активные аренды'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        rental.status = 'completed'
        rental.return_date = timezone.now()
        rental.return_condition = request.data.get('return_condition', '')
        rental.return_approved_by = request.user
        rental.save()
        
        # Обновляем статус автомобиля
        car = rental.car
        car.status = 'available'
        car.save()
        
        return Response(RentalOperatorSerializer(rental).data)

    @action(detail=True, methods=['post'])
    def reject(self, request, pk=None):
        rental = self.get_object()
        if rental.status != 'pending':
            return Response(
                {'error': 'Можно отклонять только заявки в статусе "Ожидает подтверждения"'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        rental.status = 'rejected'
        rental.rejection_reason = request.data.get('rejection_reason')
        rental.save()
        
        # Возвращаем статус автомобиля на "доступен"
        car = rental.car
        car.status = 'available'
        car.save()
        
        return Response(RentalOperatorSerializer(rental).data)

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def pay_penalty(request, pk):
    try:
        # Проверяем, что штраф принадлежит пользователю
        penalty = Penalty.objects.get(id=pk, rental__user=request.user)
        
        # Помечаем штраф как оплаченный
        penalty.is_paid = True
        penalty.paid_at = timezone.now()
        penalty.save()
        
        return Response({'status': 'success', 'message': 'Штраф успешно оплачен'})
    except Penalty.DoesNotExist:
        return Response(
            {'error': 'Штраф не найден или не принадлежит вам'},
            status=status.HTTP_404_NOT_FOUND
        )

def calculate_discount(user):
    """Рассчитывает текущую скидку пользователя на основе количества завершенных аренд в текущем месяце"""
    from django.utils import timezone
    import datetime
    
    
    # Получаем текущий месяц и год
    now = timezone.now()
    current_month = now.month
    current_year = now.year
    
    
    # Получаем все завершенные аренды пользователя
    completed_rentals = Rental.objects.filter(
        user=user,
        status='completed'
    )
    
    
    # Фильтруем аренды по текущему месяцу
    completed_rentals_this_month = []
    for rental in completed_rentals:
        if rental.return_date:
            try:
                
                # Проверяем, является ли return_date объектом datetime
                if isinstance(rental.return_date, datetime.datetime):
                    if rental.return_date.month == current_month and rental.return_date.year == current_year:
                        completed_rentals_this_month.append(rental)
                else:
                    # Если return_date - строка или другой тип, пытаемся преобразовать в datetime
                    return_date_str = str(rental.return_date).replace('Z', '+00:00')
                    
                    # Пробуем разные форматы даты
                    try:
                        return_date = timezone.datetime.fromisoformat(return_date_str)
                    except ValueError:
                        # Если формат не ISO, пробуем другие форматы
                        try:
                            return_date = datetime.datetime.strptime(return_date_str, '%Y-%m-%d %H:%M:%S%z')
                        except ValueError:
                            try:
                                return_date = datetime.datetime.strptime(return_date_str, '%Y-%m-%d %H:%M:%S')
                            except ValueError:
                                try:
                                    return_date = datetime.datetime.strptime(return_date_str, '%Y-%m-%d')
                                except ValueError:
                                    continue
                    
                    if return_date.month == current_month and return_date.year == current_year:
                        completed_rentals_this_month.append(rental)
            except (ValueError, AttributeError, TypeError):
                pass
    
    completed_rentals_count = len(completed_rentals_this_month)
    
    # Определяем уровень скидки на основе количества завершенных аренд
    discount_id = None
    if completed_rentals_count >= 20:
        discount_id = 4  # 20%
    elif completed_rentals_count >= 10:
        discount_id = 3  # 15%
    elif completed_rentals_count >= 5:
        discount_id = 2  # 10%
    elif completed_rentals_count >= 3:
        discount_id = 1  # 5%
    
    
    # Обновляем скидку пользователя непосредственно в модели User
    try:
        if discount_id:
            discount = Discount.objects.get(id=discount_id)
            user.discount = discount
            user.save()
            return discount.discount_rate
        else:
            # Если нет скидки, устанавливаем None
            user.discount = None
            user.save()
            return 0
    except Exception:
        return 0

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def create_rental(request):
    """Создание новой заявки на аренду с учетом скидки"""
    # Получаем данные из запроса
    car_id = request.data.get('car_id')
    start_date = request.data.get('start_date')
    end_date = request.data.get('end_date')
    personal_info = request.data.get('personal_info', {})
    total_price = request.data.get('total_price')
    applied_discount = request.data.get('applied_discount', 0)
    
    # Проверяем, что все необходимые данные предоставлены
    if not all([car_id, start_date, end_date, total_price]):
        return Response(
            {'error': 'Необходимо указать автомобиль, даты аренды и стоимость'},
            status=status.HTTP_400_BAD_REQUEST
        )
    
    try:
        # Получаем автомобиль
        car = Car.objects.get(id=car_id)
        
        # Проверяем, доступен ли автомобиль
        if car.status != 'available':
            return Response(
                {'error': 'Автомобиль недоступен для аренды'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Создаем объект аренды
        rental = Rental.objects.create(
            user=request.user,
            car=car,
            start_date=start_date,
            end_date=end_date,
            personal_info=personal_info,
            total_price=total_price,
            status='pending',
            applied_discount=applied_discount  # Используем скидку, переданную с фронтенда
        )
        
        # Обновляем статус автомобиля
        car.status = 'pending'
        car.save()
        
        # Возвращаем данные созданной аренды
        serializer = RentalSerializer(rental)
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    
    except Car.DoesNotExist:
        return Response(
            {'error': 'Автомобиль не найден'},
            status=status.HTTP_404_NOT_FOUND
        )
    except Exception as e:
        return Response(
            {'error': str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_user_discount(request):
    """Получение текущей скидки пользователя"""
    
    try:
        # Рассчитываем скидку на основе всех завершенных аренд
        discount_rate = calculate_discount(request.user)
        
        # Логируем информацию о скидке
        
        return Response({'discount': discount_rate})
    except Exception:
        return Response({'discount': 0})

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def debug_discount(request):
    """Отладочный эндпоинт для проверки расчета скидок"""
    from django.utils import timezone
    import datetime
    
    user = request.user
    now = timezone.now()
    current_month = now.month
    current_year = now.year
    
    
    # Получаем все завершенные аренды пользователя
    completed_rentals = Rental.objects.filter(
        user=user,
        status='completed'
    )
    
    
    # Собираем информацию о каждой аренде
    rental_info = []
    for rental in completed_rentals:
        is_in_current_month = False
        return_date_str = None
        
        if rental.return_date:
            try:
                if isinstance(rental.return_date, datetime.datetime):
                    is_in_current_month = (rental.return_date.month == current_month and 
                                          rental.return_date.year == current_year)
                    return_date_str = rental.return_date.isoformat()
                else:
                    # Если return_date - строка или другой тип, пытаемся преобразовать в datetime
                    return_date_str = str(rental.return_date)
                    return_date = timezone.datetime.fromisoformat(return_date_str.replace('Z', '+00:00'))
                    is_in_current_month = (return_date.month == current_month and 
                                          return_date.year == current_year)
            except (ValueError, AttributeError, TypeError):
                is_in_current_month = False
                
        rental_info.append({
            'id': rental.id,
            'return_date': return_date_str,
            'status': rental.status,
            'in_current_month': is_in_current_month
        })
    
    # Рассчитываем скидку
    discount = calculate_discount(user)
    
    # Получаем обновленную информацию о пользователе
    user.refresh_from_db()
    current_discount = user.discount.discount_rate if user.discount else 0
    
    # Подсчитываем количество аренд в текущем месяце
    rentals_this_month = sum(1 for r in rental_info if r['in_current_month'])
    
    # Определяем, сколько аренд нужно до следующего уровня скидки
    next_discount_level = None
    next_discount_rate = None
    rentals_needed = 0
    
    if rentals_this_month < 3:
        next_discount_level = 5
        next_discount_rate = 5
        rentals_needed = 3 - rentals_this_month
    elif rentals_this_month < 5:
        next_discount_level = 10
        next_discount_rate = 10
        rentals_needed = 5 - rentals_this_month
    elif rentals_this_month < 10:
        next_discount_level = 15
        next_discount_rate = 15
        rentals_needed = 10 - rentals_this_month
    elif rentals_this_month < 20:
        next_discount_level = 20
        next_discount_rate = 20
        rentals_needed = 20 - rentals_this_month
    
    return Response({
        'current_month': current_month,
        'current_year': current_year,
        'total_completed_rentals': completed_rentals.count(),
        'rentals_this_month': rentals_this_month,
        'rental_details': rental_info,
        'calculated_discount': discount,
        'current_user_discount': current_discount,
        'next_discount_level': next_discount_level,
        'rentals_needed_for_next_level': rentals_needed
    })

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_profile(request):
    user = request.user
    serializer = UserSerializer(user)
    return Response(serializer.data)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def debug_rental_data(request, rental_id):
    """Отладочный эндпоинт для проверки данных аренды"""
    try:
        rental = Rental.objects.get(id=rental_id)
        
        # Проверяем права доступа
        if rental.user != request.user and not request.user.is_staff:
            return Response(
                {'error': 'У вас нет прав для просмотра этой аренды'},
                status=status.HTTP_403_FORBIDDEN
            )
        
        # Собираем подробную информацию об аренде
        rental_data = {
            'id': rental.id,
            'user': rental.user.username,
            'car': {
                'id': rental.car.id,
                'brand': rental.car.brand,
                'model': rental.car.model
            },
            'start_date': rental.start_date,
            'end_date': rental.end_date,
            'return_date': rental.return_date,
            'total_price': rental.total_price,
            'personal_info': rental.personal_info,
            'status': rental.status,
            'created_at': rental.created_at,
            'approved_at': rental.approved_at,
            'return_condition': rental.return_condition,
            'applied_discount': rental.applied_discount
        }
        
        return Response(rental_data)
    
    except Rental.DoesNotExist:
        return Response(
            {'error': 'Аренда не найдена'},
            status=status.HTTP_404_NOT_FOUND
        )

@api_view(['PUT'])
@permission_classes([IsAuthenticated])
def update_profile(request):
    """Обновление профиля пользователя"""
    user = request.user
    
    print(f"Полученные данные для обновления профиля: {request.data}")
    
    # Обновляем поля пользователя
    if 'full_name' in request.data and request.data['full_name']:
        # Разбиваем полное имя на части (Имя Отчество Фамилия)
        name_parts = request.data['full_name'].split()
        user.first_name = name_parts[0] if len(name_parts) > 0 else ''
        user.middle_name = name_parts[1] if len(name_parts) > 1 else ''
        user.last_name = ' '.join(name_parts[2:]) if len(name_parts) > 2 else ''
    
    if 'email' in request.data:
        user.email = request.data['email']
    
    if 'phone' in request.data:
        user.phone = request.data['phone']
    
    if 'address' in request.data:
        user.address = request.data['address']
    
    if 'passport_number' in request.data:
        user.passport_number = request.data['passport_number']
    
    if 'driver_license' in request.data:
        user.driver_license = request.data['driver_license']
    
    user.save()
    
    # Возвращаем обновленные данные
    return get_profile(request)

class AccountingViewSet(viewsets.ViewSet):
    permission_classes = [permissions.IsAuthenticated]
    
    @action(detail=False, methods=['get'])
    def penalties(self, request):
        """Получить список штрафов с фильтрацией по статусу и периоду"""
        # Получаем параметры запроса
        status = request.query_params.get('status', None)  # paid, unpaid, all
        period = request.query_params.get('period', 'all')  # week, month, half_year, all
        
        # Базовый запрос
        penalties = Penalty.objects.all()
        
        # Фильтрация по статусу
        if status == 'paid':
            penalties = penalties.filter(is_paid=True)
        elif status == 'unpaid':
            penalties = penalties.filter(is_paid=False)
        
        # Фильтрация по периоду
        now = timezone.now()
        if period == 'week':
            start_date = now - timezone.timedelta(days=7)
            penalties = penalties.filter(created_at__gte=start_date)
        elif period == 'month':
            start_date = now - timezone.timedelta(days=30)
            penalties = penalties.filter(created_at__gte=start_date)
        elif period == 'half_year':
            start_date = now - timezone.timedelta(days=180)
            penalties = penalties.filter(created_at__gte=start_date)
        
        # Сериализуем данные
        serializer = PenaltySerializer(penalties, many=True)
        
        # Рассчитываем общую сумму оплаченных штрафов
        total_paid = penalties.filter(is_paid=True).aggregate(Sum('amount'))['amount__sum'] or 0
        
        return Response({
            'penalties': serializer.data,
            'total_paid': total_paid
        })
    
    @action(detail=False, methods=['get'])
    def statistics(self, request):
        """Получить статистику доходов и расходов"""
        # Получаем параметры запроса
        period = request.query_params.get('period', 'month')  # week, month, half_year, year, all
        include_penalties = request.query_params.get('include_penalties', 'false') == 'true'
        
        # Определяем начальную дату периода
        now = timezone.now()
        if period == 'week':
            start_date = now - timezone.timedelta(days=7)
            date_format = '%d.%m'
            delta = timezone.timedelta(days=1)
        elif period == 'month':
            start_date = now - timezone.timedelta(days=30)
            date_format = '%d.%m'
            delta = timezone.timedelta(days=1)
        elif period == 'half_year':
            start_date = now - timezone.timedelta(days=180)
            date_format = '%m.%Y'
            delta = timezone.timedelta(days=30)
        elif period == 'year':
            start_date = now - timezone.timedelta(days=365)
            date_format = '%m.%Y'
            delta = timezone.timedelta(days=30)
        else:
            start_date = now - timezone.timedelta(days=30)  # По умолчанию месяц
            date_format = '%d.%m'
            delta = timezone.timedelta(days=1)
        
        # Получаем данные о доходах (аренды)
        rentals = Rental.objects.filter(
            status='completed',
            return_date__gte=start_date
        )
        
        # Получаем данные о расходах (обслуживание)
        maintenances = Maintenance.objects.filter(
            status='completed',
            completed_date__gte=start_date
        )
        
        # Получаем данные о штрафах
        penalties = Penalty.objects.filter(
            is_paid=True,
            paid_at__gte=start_date
        ) if include_penalties else []
        
        # Формируем данные для графика
        current_date = start_date
        income_data = []
        expense_data = []
        labels = []
        
        while current_date <= now:
            # Форматируем дату для метки
            label = current_date.strftime(date_format)
            labels.append(label)
            
            # Рассчитываем доходы за период
            if period in ['week', 'month']:
                # Для недели и месяца считаем по дням
                period_start = current_date
                period_end = current_date + timezone.timedelta(days=1)
            else:
                # Для полугода и года считаем по месяцам
                period_start = current_date
                period_end = current_date + timezone.timedelta(days=30)
            
            # Доходы от аренды
            rental_income = rentals.filter(
                return_date__gte=period_start,
                return_date__lt=period_end
            ).aggregate(Sum('total_price'))['total_price__sum'] or 0
            
            # Доходы от штрафов
            penalty_income = penalties.filter(
                paid_at__gte=period_start,
                paid_at__lt=period_end
            ).aggregate(Sum('amount'))['amount__sum'] or 0 if include_penalties else 0
            
            # Расходы на обслуживание
            maintenance_expense = maintenances.filter(
                completed_date__gte=period_start,
                completed_date__lt=period_end
            ).aggregate(Sum('cost'))['cost__sum'] or 0
            
            # Добавляем данные
            income_data.append(float(rental_income) + float(penalty_income))
            expense_data.append(float(maintenance_expense))
            
            # Переходим к следующему периоду
            current_date += delta
        
        # Рассчитываем итоговые суммы
        total_income = sum(income_data)
        total_expense = sum(expense_data)
        total_profit = total_income - total_expense
        
        # Получаем данные о популярных автомобилях
        # Группируем аренды по автомобилям и считаем количество аренд для каждого автомобиля
        popular_cars_data = []
        try:
            # Получаем все завершенные аренды за указанный период
            completed_rentals = Rental.objects.filter(
                status='completed',
                return_date__gte=start_date
            )
            
            # Группируем аренды по автомобилям и считаем их количество
            car_rental_counts = {}
            for rental in completed_rentals:
                car_id = rental.car.id
                car_name = f"{rental.car.brand} {rental.car.model}"
                
                if car_id in car_rental_counts:
                    car_rental_counts[car_id]['rentals'] += 1
                else:
                    car_rental_counts[car_id] = {
                        'name': car_name,
                        'rentals': 1
                    }
            
            # Сортируем автомобили по количеству аренд (в порядке убывания)
            sorted_cars = sorted(
                car_rental_counts.values(),
                key=lambda x: x['rentals'],
                reverse=True
            )
            
            # Берем топ-5 автомобилей
            popular_cars_data = sorted_cars[:5]
        except Exception as e:
            print(f"Ошибка при получении данных о популярных автомобилях: {str(e)}")
        
        # Рассчитываем среднюю длительность аренды
        average_rental_duration = 0
        try:
            if rentals.exists():
                total_days = 0
                rental_count = 0
                for rental in rentals:
                    # Рассчитываем разницу между датой возврата и датой начала аренды
                    if rental.return_date and rental.start_date:
                        # Преобразуем оба значения к одному типу (datetime.date)
                        return_date = rental.return_date.date() if isinstance(rental.return_date, datetime.datetime) else rental.return_date
                        start_date = rental.start_date.date() if isinstance(rental.start_date, datetime.datetime) else rental.start_date
                        
                        rental_days = (return_date - start_date).days
                        if rental_days >= 0:  # Проверяем, что дата возврата не раньше даты начала
                            total_days += max(1, rental_days)  # Минимум 1 день
                            rental_count += 1
                
                if rental_count > 0:
                    average_rental_duration = round(total_days / rental_count, 1)
                else:
                    # Если нет корректных данных, используем значение по умолчанию
                    average_rental_duration = 3.0
            else:
                # Если нет аренд, используем значение по умолчанию
                average_rental_duration = 3.0
        except Exception as e:
            print(f"Ошибка при расчете средней длительности аренды: {str(e)}")
            # В случае ошибки используем значение по умолчанию
            average_rental_duration = 3.0
        
        # Рассчитываем загрузку автопарка
        fleet_utilization = 0
        try:
            # Получаем общее количество автомобилей
            total_cars = Car.objects.count()
            
            # Получаем количество автомобилей в аренде
            rented_cars = Car.objects.filter(
                Q(status='rented') | Q(status='in_rent')
            ).count()
            
            # Рассчитываем процент загрузки
            if total_cars > 0:
                fleet_utilization = round((rented_cars / total_cars) * 100)
        except Exception as e:
            print(f"Ошибка при расчете загрузки автопарка: {str(e)}")
        
        # Рассчитываем общие затраты на обслуживание
        total_maintenance_costs = 0
        try:
            # Получаем все завершенные обслуживания
            all_maintenances = Maintenance.objects.filter(status='completed')
            
            # Суммируем затраты
            total_maintenance_costs = all_maintenances.aggregate(Sum('cost'))['cost__sum'] or 0
        except Exception as e:
            print(f"Ошибка при расчете общих затрат на обслуживание: {str(e)}")
        
        return Response({
            'labels': labels,
            'income_data': income_data,
            'expense_data': expense_data,
            'total_income': total_income,
            'total_expense': total_expense,
            'total_profit': total_profit,
            'popular_cars': popular_cars_data,
            'total_rentals': rentals.count(),
            'average_rental_duration': average_rental_duration,
            'fleet_utilization': fleet_utilization,
            'total_maintenance_costs': float(total_maintenance_costs)
        })
    
    @action(detail=False, methods=['get'])
    def tax_report(self, request):
        """Сформировать налоговый отчет"""
        try:
            # Получаем параметры запроса
            period = request.query_params.get('period', 'month')  # month, quarter, year
            
            # Определяем начальную и конечную даты периода
            now = timezone.now()
            if period == 'month':
                start_date = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
                if start_date.month == 12:
                    end_date = start_date.replace(year=start_date.year + 1, month=1) - timezone.timedelta(days=1)
                else:
                    end_date = start_date.replace(month=start_date.month + 1) - timezone.timedelta(days=1)
                
                # Русские названия месяцев
                month_names_ru = {
                    1: 'Январь', 2: 'Февраль', 3: 'Март', 4: 'Апрель', 5: 'Май', 6: 'Июнь',
                    7: 'Июль', 8: 'Август', 9: 'Сентябрь', 10: 'Октябрь', 11: 'Ноябрь', 12: 'Декабрь'
                }
                period_name = f"за {month_names_ru[start_date.month]} {start_date.year}"
            elif period == 'quarter':
                quarter = (now.month - 1) // 3 + 1
                start_date = now.replace(month=(quarter-1)*3+1, day=1, hour=0, minute=0, second=0, microsecond=0)
                if quarter == 4:
                    end_date = now.replace(year=now.year+1, month=1, day=1) - timezone.timedelta(days=1)
                else:
                    end_date = now.replace(month=quarter*3+1, day=1) - timezone.timedelta(days=1)
                period_name = f"за {quarter} квартал {now.year}"
            elif period == 'year':
                start_date = now.replace(month=1, day=1, hour=0, minute=0, second=0, microsecond=0)
                end_date = now.replace(month=12, day=31, hour=23, minute=59, second=59)
                period_name = f"за {now.year} год"
            else:
                # По умолчанию - текущий месяц
                start_date = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
                if start_date.month == 12:
                    end_date = start_date.replace(year=start_date.year + 1, month=1) - timezone.timedelta(days=1)
                else:
                    end_date = start_date.replace(month=start_date.month + 1) - timezone.timedelta(days=1)
                period_name = f"за {start_date.strftime('%B %Y')}"
            
            # Получаем данные о доходах (аренды)
            rentals = Rental.objects.filter(
                status='completed',
                return_date__gte=start_date,
                return_date__lte=end_date
            )
            rental_income = rentals.aggregate(Sum('total_price'))['total_price__sum'] or 0
            
            # Получаем данные о доходах от штрафов
            penalties = Penalty.objects.filter(
                is_paid=True,
                paid_at__gte=start_date,
                paid_at__lte=end_date
            )
            penalty_income = penalties.aggregate(Sum('amount'))['amount__sum'] or 0
            
            # Получаем данные о расходах (обслуживание)
            maintenances = Maintenance.objects.filter(
                status='completed',
                completed_date__gte=start_date,
                completed_date__lte=end_date
            )
            maintenance_expense = maintenances.aggregate(Sum('cost'))['cost__sum'] or 0
            
            # Рассчитываем итоговые суммы
            total_income = float(rental_income) + float(penalty_income)
            total_expense = float(maintenance_expense)
            total_profit = total_income - total_expense
            
            # Рассчитываем налог (условно 20% от прибыли)
            tax_amount = total_profit * 0.2
            
            # Создаем документ Word
            doc = Document()
            
            # Устанавливаем шрифт Times New Roman 14pt для всего документа
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(14)
            
            # Добавляем заголовок с форматированием 16pt полужирный черный
            heading = doc.add_paragraph()
            run = heading.add_run(f'Налоговый отчет {period_name}')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Добавляем информацию о компании
            doc.add_paragraph('ООО "Sewxrr RentCar"')
            doc.add_paragraph(f'ИНН: 1234567890')
            doc.add_paragraph(f'КПП: 123456789')
            doc.add_paragraph(f'Дата составления: {now.strftime("%d.%m.%Y")}')
            
            # Добавляем заголовок для финансовых показателей
            fin_heading = doc.add_paragraph()
            fin_run = fin_heading.add_run('Финансовые показатели')
            fin_run.font.name = 'Times New Roman'
            fin_run.font.size = Pt(16)
            fin_run.font.bold = True
            fin_run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Добавляем таблицу с финансовыми показателями
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Форматируем заголовки таблицы
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Показатель'
            hdr_cells[1].text = 'Сумма (руб.)'
            
            # Применяем форматирование к заголовкам таблицы
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)
                        run.font.bold = True
            
            # Доходы от аренды
            row_cells = table.add_row().cells
            row_cells[0].text = 'Доходы от аренды'
            row_cells[1].text = f'{rental_income:.2f}'
            
            # Доходы от штрафов
            row_cells = table.add_row().cells
            row_cells[0].text = 'Доходы от штрафов'
            row_cells[1].text = f'{penalty_income:.2f}'
            
            # Общий доход
            row_cells = table.add_row().cells
            row_cells[0].text = 'Общий доход'
            row_cells[1].text = f'{total_income:.2f}'
            
            # Расходы на обслуживание
            row_cells = table.add_row().cells
            row_cells[0].text = 'Расходы на обслуживание'
            row_cells[1].text = f'{maintenance_expense:.2f}'
            
            # Общие расходы
            row_cells = table.add_row().cells
            row_cells[0].text = 'Общие расходы'
            row_cells[1].text = f'{total_expense:.2f}'
            
            # Прибыль
            row_cells = table.add_row().cells
            row_cells[0].text = 'Прибыль'
            row_cells[1].text = f'{total_profit:.2f}'
            
            # Налог
            row_cells = table.add_row().cells
            row_cells[0].text = 'Налог (20%)'
            row_cells[1].text = f'{tax_amount:.2f}'
            
            # Чистая прибыль
            row_cells = table.add_row().cells
            row_cells[0].text = 'Чистая прибыль'
            row_cells[1].text = f'{(total_profit - tax_amount):.2f}'
            
            # Добавляем детализацию по арендам
            doc.add_paragraph()
            doc.add_paragraph()
            
            rentals_heading = doc.add_paragraph()
            rentals_run = rentals_heading.add_run('Детализация по арендам')
            rentals_run.font.name = 'Times New Roman'
            rentals_run.font.size = Pt(16)
            rentals_run.font.bold = True
            rentals_run.font.color.rgb = RGBColor(0, 0, 0)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ID'
            hdr_cells[1].text = 'Автомобиль'
            hdr_cells[2].text = 'Дата возврата'
            hdr_cells[3].text = 'Сумма (руб.)'
            
            # Применяем форматирование к заголовкам таблицы
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)
                        run.font.bold = True
            
            for rental in rentals:
                row_cells = table.add_row().cells
                row_cells[0].text = str(rental.id)
                row_cells[1].text = f'{rental.car.brand} {rental.car.model}'
                row_cells[2].text = rental.return_date.strftime('%d.%m.%Y')
                row_cells[3].text = f'{rental.total_price:.2f}'
            
            # Добавляем детализацию по штрафам
            doc.add_paragraph()
            doc.add_paragraph()
            
            penalties_heading = doc.add_paragraph()
            penalties_run = penalties_heading.add_run('Детализация по штрафам')
            penalties_run.font.name = 'Times New Roman'
            penalties_run.font.size = Pt(16)
            penalties_run.font.bold = True
            penalties_run.font.color.rgb = RGBColor(0, 0, 0)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ID'
            hdr_cells[1].text = 'Описание'
            hdr_cells[2].text = 'Дата оплаты'
            hdr_cells[3].text = 'Сумма (руб.)'
            
            # Применяем форматирование к заголовкам таблицы
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)
                        run.font.bold = True
            
            for penalty in penalties:
                row_cells = table.add_row().cells
                row_cells[0].text = str(penalty.id)
                row_cells[1].text = penalty.description
                row_cells[2].text = penalty.paid_at.strftime('%d.%m.%Y')
                row_cells[3].text = f'{penalty.amount:.2f}'
            
            # Добавляем детализацию по обслуживанию
            doc.add_paragraph()
            doc.add_paragraph()
            
            maint_heading = doc.add_paragraph()
            maint_run = maint_heading.add_run('Детализация по обслуживанию')
            maint_run.font.name = 'Times New Roman'
            maint_run.font.size = Pt(16)
            maint_run.font.bold = True
            maint_run.font.color.rgb = RGBColor(0, 0, 0)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ID'
            hdr_cells[1].text = 'Автомобиль'
            hdr_cells[2].text = 'Дата завершения'
            hdr_cells[3].text = 'Стоимость (руб.)'
            
            # Применяем форматирование к заголовкам таблицы
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)
                        run.font.bold = True
            
            for maintenance in maintenances:
                row_cells = table.add_row().cells
                row_cells[0].text = str(maintenance.id)
                row_cells[1].text = f'{maintenance.car.brand} {maintenance.car.model}'
                row_cells[2].text = maintenance.completed_date.strftime('%d.%m.%Y')
                row_cells[3].text = f'{maintenance.cost:.2f}'
            
            # Добавляем подпись
            doc.add_paragraph('\n\nПодпись руководителя: ________________')
            doc.add_paragraph('Подпись главного бухгалтера: ________________')
            
            # Сохраняем документ в буфер
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Создаем HTTP-ответ с документом
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
            # Добавляем заголовки CORS
            response['Access-Control-Allow-Origin'] = '*'
            response['Access-Control-Allow-Methods'] = 'GET, OPTIONS'
            response['Access-Control-Allow-Headers'] = 'Content-Type, Authorization'
            
            # Устанавливаем имя файла на русском
            month_names = {
                1: 'Январь', 2: 'Февраль', 3: 'Март', 4: 'Апрель', 5: 'Май', 6: 'Июнь',
                7: 'Июль', 8: 'Август', 9: 'Сентябрь', 10: 'Октябрь', 11: 'Ноябрь', 12: 'Декабрь'
            }
            
            if period == 'month':
                ru_period = f"за {month_names[start_date.month].lower()} {start_date.year}"
            elif period == 'quarter':
                quarter = (now.month - 1) // 3 + 1
                ru_period = f"за {quarter} квартал {now.year}"
            else:
                ru_period = f"за {now.year} год"
            
            filename = f'Налоговый отчет {ru_period}.docx'
            
            # Используем ASCII имя файла для обычного параметра filename
            ascii_filename = f'tax_report_{period}.docx'
            
            # Используем RFC5987 для кодирования имени файла с русскими символами
            encoded_filename = quote(filename.encode('utf-8'))
            
            # Добавляем оба заголовка для максимальной совместимости
            response['Content-Disposition'] = f'attachment; filename="{ascii_filename}"; filename*=UTF-8\'\'{encoded_filename}'
            
            return response
        except Exception as e:
            print(f"Error generating tax report: {str(e)}")
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['GET'])
@permission_classes([AllowAny])  # Временно разрешаем доступ всем для тестирования
def car_financial_history(request):
    """Получить историю доходов и расходов по каждой машине"""
    try:
        # Получаем все машины
        cars = Car.objects.all()
        
        # Создаем словарь для хранения финансовой информации по каждой машине
        car_finances = {}
        
        # Для каждой машины получаем историю аренд и обслуживаний
        for car in cars:
            car_id = car.id
            
            # Получаем завершенные аренды для этой машины
            rentals = Rental.objects.filter(
                car=car,
                status='completed'
            )
            
            # Рассчитываем общий доход от аренд
            total_income = rentals.aggregate(Sum('total_price'))['total_price__sum'] or 0
            
            # Получаем завершенные обслуживания для этой машины
            maintenances = Maintenance.objects.filter(
                car=car,
                status='completed'
            )
            
            # Рассчитываем общие расходы на обслуживание
            total_expenses = maintenances.aggregate(Sum('cost'))['cost__sum'] or 0
            
            # Рассчитываем эффективность (прибыльность)
            efficiency = 0
            if total_income > 0:
                efficiency = round((total_income - total_expenses) / total_income * 100)
            
            # Сохраняем финансовую информацию для этой машины
            car_finances[car_id] = {
                'id': car_id,
                'brand': car.brand,
                'model': car.model,
                'total_income': float(total_income),
                'total_expenses': float(total_expenses),
                'efficiency': efficiency
            }
        
        # Преобразуем словарь в список для ответа API
        result = list(car_finances.values())
        
        return Response(result)
    except Exception as e:
        print(f"Ошибка при получении финансовой истории: {str(e)}")
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def debug_rentals(request):
    """Отладочный эндпоинт для проверки статуса аренд и их дат"""
    import datetime
    from django.utils import timezone
    
    user = request.user
    now = timezone.now()
    current_month = now.month
    current_year = now.year
    
    # Получаем все аренды пользователя
    all_rentals = Rental.objects.filter(user=user)
    
    # Собираем информацию о каждой аренде
    rental_info = []
    for rental in all_rentals:
        return_date_str = None
        return_date_obj = None
        is_in_current_month = False
        
        if rental.return_date:
            try:
                if isinstance(rental.return_date, datetime.datetime):
                    return_date_str = rental.return_date.isoformat()
                    return_date_obj = rental.return_date
                    is_in_current_month = (rental.return_date.month == current_month and 
                                          rental.return_date.year == current_year)
                else:
                    # Если return_date - строка или другой тип, пытаемся преобразовать в datetime
                    return_date_str = str(rental.return_date)
                    try:
                        return_date_obj = timezone.datetime.fromisoformat(return_date_str.replace('Z', '+00:00'))
                    except ValueError:
                        try:
                            return_date_obj = datetime.datetime.strptime(return_date_str, '%Y-%m-%d %H:%M:%S%z')
                        except ValueError:
                            try:
                                return_date_obj = datetime.datetime.strptime(return_date_str, '%Y-%m-%d %H:%M:%S')
                            except ValueError:
                                try:
                                    return_date_obj = datetime.datetime.strptime(return_date_str, '%Y-%m-%d')
                                except ValueError:
                                    return_date_obj = None
                    
                    if return_date_obj:
                        is_in_current_month = (return_date_obj.month == current_month and 
                                              return_date_obj.year == current_year)
            except Exception:
                pass
        
        rental_info.append({
            'id': rental.id,
            'status': rental.status,
            'start_date': str(rental.start_date),
            'end_date': str(rental.end_date),
            'return_date': return_date_str,
            'return_date_type': type(rental.return_date).__name__,
            'is_in_current_month': is_in_current_month,
            'car': str(rental.car)
        })
    
    # Группируем аренды по статусу
    status_counts = {}
    for rental in all_rentals:
        status = rental.status
        if status in status_counts:
            status_counts[status] += 1
        else:
            status_counts[status] = 1
    
    # Подсчитываем количество завершенных аренд в текущем месяце
    completed_in_current_month = sum(1 for r in rental_info 
                                    if r['status'] == 'completed' and r['is_in_current_month'])
    
    return Response({
        'current_month': current_month,
        'current_year': current_year,
        'total_rentals': all_rentals.count(),
        'status_counts': status_counts,
        'completed_in_current_month': completed_in_current_month,
        'rental_details': rental_info
    })
